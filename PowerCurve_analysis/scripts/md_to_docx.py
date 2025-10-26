"""
Convert markdown power curve analysis to Word document using template.
Handles image insertion, headings, tables, formatting, and LaTeX equations.
"""

import re
import shutil
from pathlib import Path
from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# Optional LaTeX support
try:
    from lxml import etree
    from latex2mathml.converter import convert as latex_to_mathml
    LATEX_SUPPORT = True
except ImportError:
    LATEX_SUPPORT = False
    print("WARNING: LaTeX support not available (latex2mathml not installed)")

# Monkey patch Document() to accept .dotx templates
import docx.api
from docx.package import Package
_original_document = docx.api.Document

def patched_document(docx_path=None):
    """Patched Document loader that accepts both .docx and .dotx files"""
    if docx_path is None:
        return _original_document()

    # Try loading normally
    try:
        return _original_document(docx_path)
    except ValueError as e:
        # If it's a template content type error, use blank document
        # Template styles will be applied when we save
        if 'template.main+xml' in str(e):
            print(f"WARNING: Template detected (.dotx), creating document with default styles")
            print(f"   (Template formatting may not be preserved)")
            return _original_document()  # Create blank document
        raise

docx.api.Document = patched_document
Document = patched_document

# Path to Microsoft's MathML to OMML converter
MML2OMML_XSL_PATH = "/mnt/c/Program Files/Microsoft Office/root/Office16/MML2OMML.XSL"

def latex_to_omml(latex_string):
    """
    Convert LaTeX math string to Office Math Markup Language (OMML).

    Args:
        latex_string: LaTeX math expression (without $ delimiters)

    Returns:
        lxml Element containing OMML representation
    """
    if not LATEX_SUPPORT:
        print(f"WARNING: LaTeX conversion skipped (not installed): {latex_string[:50]}...")
        return None

    try:
        # Convert LaTeX to MathML
        mathml_string = latex_to_mathml(latex_string)

        # Parse MathML XML
        mathml_tree = etree.fromstring(mathml_string.encode('utf-8'))

        # Load Microsoft's MathML to OMML XSLT converter
        xslt = etree.parse(MML2OMML_XSL_PATH)
        transform = etree.XSLT(xslt)

        # Convert MathML to OMML
        omml_tree = transform(mathml_tree)

        return omml_tree.getroot()

    except Exception as e:
        print(f"WARNING: LaTeX conversion error: {latex_string[:50]}... -> {e}")
        return None

def parse_markdown_to_docx(md_path, template_path, output_path, figures_dir):
    """
    Convert markdown file to Word document using template.

    Args:
        md_path: Path to markdown file
        template_path: Path to .dotx template file
        output_path: Path for output .docx file
        figures_dir: Directory containing figure images
    """
    # Load template directly (monkey patch handles .dotx)
    doc = Document(str(template_path))

    # Read markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Track if we're in a table
    in_table = False
    table_lines = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Handle headers (strip markdown numbering like "1.", "1.1", "1.1.1")
        if line.startswith('# '):
            heading_text = line[2:].strip()
            # Remove numbering pattern at start (e.g., "1. ", "1.1 ", "1.1.1 ")
            heading_text = re.sub(r'^\d+(\.\d+)*\.?\s+', '', heading_text)
            heading = doc.add_heading(heading_text, level=1)
            # Disable numbering for appendix headings
            if heading_text.startswith('Appendix') or heading_text.startswith('APPENDIX'):
                pPr = heading._element.get_or_add_pPr()
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    pPr.remove(numPr)

        elif line.startswith('## '):
            heading_text = line[3:].strip()
            heading_text = re.sub(r'^\d+(\.\d+)*\.?\s+', '', heading_text)
            heading = doc.add_heading(heading_text, level=2)
            # Disable numbering for appendix headings
            if heading_text.startswith('Appendix') or heading_text.startswith('APPENDIX'):
                pPr = heading._element.get_or_add_pPr()
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    pPr.remove(numPr)

        elif line.startswith('### '):
            heading_text = line[4:].strip()
            heading_text = re.sub(r'^\d+(\.\d+)*\.?\s+', '', heading_text)
            heading = doc.add_heading(heading_text, level=3)
            # Disable numbering for appendix headings
            if heading_text.startswith('Appendix') or heading_text.startswith('APPENDIX'):
                pPr = heading._element.get_or_add_pPr()
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    pPr.remove(numPr)

        elif line.startswith('#### '):
            heading_text = line[5:].strip()
            heading_text = re.sub(r'^\d+(\.\d+)*\.?\s+', '', heading_text)
            heading = doc.add_heading(heading_text, level=4)
            # Disable numbering for appendix headings
            if heading_text.startswith('Appendix') or heading_text.startswith('APPENDIX'):
                pPr = heading._element.get_or_add_pPr()
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    pPr.remove(numPr)

        # Handle images: ![caption](path)
        elif line.strip().startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
            if match:
                caption, img_filename = match.groups()

                # Try multiple locations for the image
                img_path = None

                # First try: figures directory
                candidate = Path(figures_dir) / img_filename
                if candidate.exists():
                    img_path = candidate

                # Second try: parent directory (for images without figures/ prefix)
                if img_path is None:
                    candidate = Path(figures_dir).parent / img_filename
                    if candidate.exists():
                        img_path = candidate

                # Third try: assume it's already a full path relative to figures_dir
                if img_path is None:
                    # Remove 'figures/' prefix if present
                    cleaned_name = img_filename.replace('figures/', '')
                    candidate = Path(figures_dir) / cleaned_name
                    if candidate.exists():
                        img_path = candidate

                # Check if file was found
                if img_path and img_path.exists():
                    # Add image
                    doc.add_picture(str(img_path), width=Inches(6.5))

                    # Add caption if provided
                    if caption:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(caption)
                        run.font.italic = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(64, 64, 64)
                else:
                    # Image not found - add placeholder text
                    p = doc.add_paragraph()
                    run = p.add_run(f"[Image not found: {img_filename}]")
                    run.font.color.rgb = RGBColor(255, 0, 0)

        # Handle tables (markdown format)
        elif '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)

            # Check if next line is end of table
            if i + 1 >= len(lines) or '|' not in lines[i + 1]:
                # Process complete table
                process_table(doc, table_lines)
                in_table = False
                table_lines = []

        # Handle bold **text**
        elif '**' in line:
            p = doc.add_paragraph()
            process_inline_formatting(p, line)

        # Handle bullet lists
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            try:
                p = doc.add_paragraph(text, style='List Bullet')
            except KeyError:
                # Style doesn't exist in template, use plain paragraph with bullet
                p = doc.add_paragraph(text)
                p.style = 'Normal'

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            try:
                p = doc.add_paragraph(text, style='List Number')
            except KeyError:
                # Style doesn't exist in template, use plain paragraph
                p = doc.add_paragraph(text)
                p.style = 'Normal'

        # Handle code blocks ```
        elif line.strip().startswith('```'):
            # Skip opening ```
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1

            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)

        # Regular paragraph
        elif line.strip():
            p = doc.add_paragraph()
            process_inline_formatting(p, line)

        # Empty line - skip
        else:
            pass

        i += 1

    # Save document
    doc.save(output_path)
    print(f"[SUCCESS] Document saved to: {output_path}")

def process_inline_formatting(paragraph, text):
    """
    Process inline formatting like **bold**, *italic*, `code`, and $latex$.
    """
    # First split on LaTeX inline math $...$
    parts = re.split(r'(\$[^\$]+\$)', text)

    for part in parts:
        if part.startswith('$') and part.endswith('$') and len(part) > 2:
            # Inline LaTeX equation
            latex_code = part[1:-1]  # Remove $ delimiters
            omml_element = latex_to_omml(latex_code)

            if omml_element is not None:
                # Insert OMML equation into paragraph
                paragraph._element.append(omml_element)
            else:
                # Fallback: insert as text if conversion failed
                run = paragraph.add_run(part)
                run.font.color.rgb = RGBColor(255, 0, 0)

        else:
            # Process other inline formatting
            process_inline_formatting_no_latex(paragraph, part)


def process_inline_formatting_no_latex(paragraph, text):
    """
    Process inline formatting like **bold**, *italic*, `code` (no LaTeX).
    """
    # Pattern for bold **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Code text
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            # Regular text
            paragraph.add_run(part)

def process_table(doc, table_lines):
    """
    Process markdown table and add to document.
    """
    # Filter out separator line (with dashes)
    data_lines = [line for line in table_lines if not re.match(r'^\|[\s\-:|]+\|$', line)]

    if not data_lines:
        return

    # Parse table rows
    rows = []
    for line in data_lines:
        # Split by | and clean
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first/last elements from split
        cells = [c for c in cells if c]
        rows.append(cells)

    if not rows:
        return

    # Create Word table
    num_rows = len(rows)
    num_cols = len(rows[0])

    table = doc.add_table(rows=num_rows, cols=num_cols)
    try:
        # Try with dash and spaces (Word's standard naming)
        table.style = 'List Table 3 - Accent 5'
    except KeyError:
        # Try without dash
        try:
            table.style = 'List Table 3 Accent 5'
        except KeyError:
            # Try alternate naming
            try:
                table.style = 'Light List - Accent 5'
            except KeyError:
                try:
                    table.style = 'Table Grid'
                except KeyError:
                    pass  # Use default

    # Populate table
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text

            # Header row formatting
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

def main():
    """Main conversion function."""
    # Define paths
    base_dir = Path("/mnt/c/Users/klaus/klauspython/Latam")

    md_file = base_dir / "PowerCurve_analysis" / "power_curve_analysis.md"
    template_file = base_dir / "latam_hybrid" / "docs" / "WFD template.docx"
    output_file = base_dir / "PowerCurve_analysis" / "power_curve_analysis.docx"
    figures_dir = base_dir / "PowerCurve_analysis" / "figures"

    # Verify files exist
    if not md_file.exists():
        print(f"[ERROR] Markdown file not found: {md_file}")
        return

    if not template_file.exists():
        print(f"[ERROR] Template file not found: {template_file}")
        return

    if not figures_dir.exists():
        print(f"[WARNING] Figures directory not found: {figures_dir}")

    print(f"Converting: {md_file.name}")
    print(f"Template: {template_file.name}")
    print(f"Figures: {figures_dir}")
    print(f"Output: {output_file}")
    print()

    # Convert
    parse_markdown_to_docx(
        md_path=md_file,
        template_path=template_file,
        output_path=output_file,
        figures_dir=figures_dir
    )

    print()
    print("[SUCCESS] Conversion complete!")
    print(f"Open: {output_file}")

if __name__ == "__main__":
    main()
